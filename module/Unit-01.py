from docx.oxml.ns import qn  #用于设置中日文字体

from docx.shared import Pt
from docx import Document

# 从文件创建文档对象
document = Document()

with open("天声人语.md", 'r', encoding='UTF-8') as note_file:
    cards = []
    lines = note_file.readlines()
    i = 0

    print(len(lines))
    while i < len(lines):

        if i % 3 == 0:  # 添加段落,设置字体为明朝体
            #RightColumn= "RightColumnS" + str(i)
            RightColumn = document.add_paragraph()
            Text = RightColumn.add_run(lines[i].replace('\n', ''))
            Text.font.name = 'MS Mincho'
            Text._element.rPr.rFonts.set(qn('w:eastAsia'), u'MS Mincho')
            RightColumn.paragraph_format.first_line_indent = Pt(11)
            i = i + 1
        elif i % 3 == 1:
            #LeftColumn= "LeftColumn" + str(i)
            LeftColumn = document.add_paragraph()
            Text = LeftColumn.add_run(lines[i].replace('\n', ''))
            Text.font.name = '华文楷体'
            Text._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
            LeftColumn.paragraph_format.first_line_indent = Pt(22)
            i = i + 1
        else:
            i = i + 1
            continue
    ## 测试是否可以将自带的排版格式用于表格
    # table = document.add_table(rows=6, cols=2)
# 保存文档
document.save('天声人语.docx')
