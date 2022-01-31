'''
哈哈，恭喜你发现了一个小彩蛋，这个脚本是YiPai的前身，也是NoHeartPen的黑历史
'''

from docx.oxml.ns import qn  # 用于设置中日文字体
from docx.shared import Pt
from docx import Document

# 从文件创建文档对象
document = Document()

with open("天声人语.md", 'r', encoding='UTF-8') as note_file:
    cards = []
    lines = note_file.readlines()
    table = document.add_table(rows=12, cols=2)
    row = 0
    clo = 1

    cell = table.cell(row, clo)
    # 添加段落,设置字体为明朝体
    p1 = cell.add_paragraph()
    r1 = p1.add_run(lines[0].replace('\n', ''))
    r1.font.name = 'MS Mincho'
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), u'MS Mincho')
    p1.paragraph_format.first_line_indent = Pt(11)
    row = row + 1

    cell = table.cell(row, clo)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(lines[1].replace('\n', ''))
    r2.font.name = '华文楷体'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
    p2.paragraph_format.first_line_indent = Pt(22)
    row = row + 1

    cell = table.cell(row, clo)
    p3 = cell.add_paragraph()
    r3 = p3.add_run(lines[3].replace('\n', ''))
    r3.font.name = 'MS Mincho'
    r3._element.rPr.rFonts.set(qn('w:eastAsia'), u'MS Mincho')
    p3.paragraph_format.first_line_indent = Pt(11)
    row = row + 1

    cell = table.cell(row, clo)
    p4 = cell.add_paragraph()
    r4 = p4.add_run(lines[4].replace('\n', ''))
    r4.font.name = '华文楷体'
    r4._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
    p4.paragraph_format.first_line_indent = Pt(22)
    row = row + 1

    cell = table.cell(row, clo)
    p5 = cell.add_paragraph()
    r5 = p5.add_run(lines[6].replace('\n', ''))
    r5.font.name = 'MS Mincho'
    r5._element.rPr.rFonts.set(qn('w:eastAsia'), u'MS Mincho')
    p5.paragraph_format.first_line_indent = Pt(11)
    row = row + 1

    cell = table.cell(row, clo)
    p6 = cell.add_paragraph()
    r6 = p6.add_run(lines[7].replace('\n', ''))
    r6.font.name = '华文楷体'
    r6._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
    p6.paragraph_format.first_line_indent = Pt(22)
    row = row + 1

    cell = table.cell(row, clo)
    p7 = cell.add_paragraph()
    r7 = p7.add_run(lines[9].replace('\n', ''))
    r7.font.name = 'MS Mincho'
    r7._element.rPr.rFonts.set(qn('w:eastAsia'), u'MS Mincho')
    p7.paragraph_format.first_line_indent = Pt(11)
    row = row + 1

    cell = table.cell(row, clo)
    p8 = cell.add_paragraph()
    r8 = p8.add_run(lines[10].replace('\n', ''))
    r8.font.name = '华文楷体'
    r8._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
    p8.paragraph_format.first_line_indent = Pt(22)
    row = row + 1

    cell = table.cell(row, clo)
    p9 = cell.add_paragraph()
    r9 = p9.add_run(lines[12].replace('\n', ''))
    r9.font.name = 'MS Mincho'
    r9._element.rPr.rFonts.set(qn('w:eastAsia'), u'MS Mincho')
    p9.paragraph_format.first_line_indent = Pt(11)
    row = row + 1

    cell = table.cell(row, clo)
    p10 = cell.add_paragraph()
    r10 = p10.add_run(lines[13].replace('\n', ''))
    r10.font.name = '华文楷体'
    r10._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
    p10.paragraph_format.first_line_indent = Pt(22)
    row = row + 1

    cell = table.cell(row, clo)
    p11 = cell.add_paragraph()
    r11 = p11.add_run(lines[15].replace('\n', ''))
    r11.font.name = 'MS Mincho'
    r11._element.rPr.rFonts.set(qn('w:eastAsia'), u'MS Mincho')
    p11.paragraph_format.first_line_indent = Pt(11)
    row = row + 1

    cell = table.cell(row, clo)
    p12 = cell.add_paragraph()
    r12 = p12.add_run(lines[16].replace('\n', ''))
    r12.font.name = '华文楷体'
    r12._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
    p12.paragraph_format.first_line_indent = Pt(22)

# 保存文档
document.save('天声人语.docx')
