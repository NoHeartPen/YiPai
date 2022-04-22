from docx import Document
from docx.oxml.ns import qn  # 用于设置中日文字体
from docx.shared import Pt  # 设置文本缩进
import os
from pathlib import Path

# Main
Inputpath = os.getcwd()
p = Path(Inputpath)


# module\target\md2docx_1lang.py
def md2docx_2lang(file):
    document = Document()
    with open(file, 'r', encoding='UTF-8') as ProcessFile:
        lines = ProcessFile.readlines()
        i = 0
        row = 0
        table = document.add_table(rows=int(len(lines) / 2) + 1, cols=2)

        while i < len(lines):
            if i % 2 == 0:
                Rightcell = table.cell(row, 0)
                Rightcell.text = lines[i].replace('\n', '')
                i = i + 1
            elif i % 2 == 1:
                LeftCell = table.cell(row, 1)
                LeftCell.text = lines[i].replace('\n', '')
                i = i + 1
                row += 1
            else:
                i = i + 1
                continue

        for row in table.rows:  #设置中文字体
            for paragraph in row.cells[1].paragraphs:
                paragraph.paragraph_format.first_line_indent = Pt(18)
                for run in paragraph.runs:
                    run.font.name = '华文楷体'
                    run.font.size = Pt(9)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')

            for paragraph in row.cells[0].paragraphs:  # 设置日文字体
                paragraph.paragraph_format.first_line_indent = Pt(9)
                for run in paragraph.runs:
                    run.font.name = 'MS Mincho'
                    run.font.size = Pt(9)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'MS Mincho')

    ProcessFile.close()
    document.save(".YiPai\.process\{}.docx".format(
        os.path.basename(file).replace(".md", "")))


def Target2lang():
    ProcessFilePath = Inputpath + '\\.YiPai\\' + '.process' + '\\'
    p = Path(ProcessFilePath)
    FileList = list(p.glob("**/*.md"))
    for file in FileList:
        md2docx_2lang(file)


Target2lang()
