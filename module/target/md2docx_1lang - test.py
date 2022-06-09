from docx import Document
from docx.oxml.ns import qn  # 用于设置中日文字体
from docx.shared import Pt  # 设置文本缩进
import os
from pathlib import Path
import json

# Main
Inputpath = os.getcwd()
p = Path(Inputpath)
ProcessFilePath = Inputpath + '\\.YiPai' + '\\'
SettingFile = ProcessFilePath + 'defaultsetting.json' # 注意配置文件没有放在.process

SettingDict = dict()

with open(SettingFile, 'r') as file:
    SettingDict = json.load(file)


# module\target\md2docx_1lang.py
def md2docx_1lang(file):
    document = Document()
    with open(file, 'r', encoding='UTF-8') as ProcessFile:
        lines = ProcessFile.readlines()
        i = 0
        row = 0
        table = document.add_table(rows=len(lines), cols=2)

        while i < len(lines):
            LeftCell = table.cell(row, 0)
            LeftCell.text = lines[i].replace('\n', '')
            i = i + 1
            row += 1

        for row in table.rows:
            for paragraph in row.cells[0].paragraphs:
                paragraph.paragraph_format.first_line_indent = Pt(
                    SettingDict['LeftFontSize'])
                for run in paragraph.runs:
                    run.font.name = SettingDict['LeftFontName']
                    run.font.size = Pt(SettingDict['LeftFontSize'])
                    run._element.rPr.rFonts.set(qn('w:eastAsia'),
                                                SettingDict['LeftFontName'])
    ProcessFile.close()
    document.save(".YiPai\.process\{}.docx".format(
        os.path.basename(file).replace(".md", "")))


# Target
def Target1lang():
    p = Path(ProcessFilePath)
    FileList = list(p.glob("**/*.md"))
    for file in FileList:
        md2docx_1lang(file)


Target1lang()
