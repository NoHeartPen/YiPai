from pathlib import Path
from docx import Document
import os
import re
from docx.oxml.ns import qn  # 用于设置中日文字体
from docx.shared import Pt  # 设置文本缩进
import shutil  # pip install shutilwhich
from docx2pdf import convert
import tkinter as tk
import ctypes # Window与C++ 默认
from tkinter import ttk

# Main
Inputpath = os.getcwd()
p = Path(Inputpath)
ProcessFilePath = Inputpath + '\\.YiPai\\.process' + '\\'


# module\situation\docx2md.py
def docx2md(docx2md_File):
    file = Document(docx2md_File)
    paragraphs = []

    for p in file.paragraphs:  # 读取所有行
        paragraphs.append(p.text + "\n")

    with open("{}.md".format(
            ProcessFilePath +
            os.path.basename(docx2md_File).replace(".docx", "")),
              'w',
              encoding='UTF-8') as OutPutFile:
        OutPutFile.writelines(paragraphs)


# module\situation\md2md.py
def md2md(InputFile):
    InputFileName = os.path.basename(InputFile)
    ProcessFileName = ProcessFilePath + InputFileName
    shutil.copy(InputFileName, ProcessFileName)


# module\target\md2docx_2lang.py
def md2docx_2lang(file):
    document = Document()
    with open(file, 'r', encoding='UTF-8') as ProcessFile:
        lines = ProcessFile.readlines()
        i = 0
        row = 0
        table = document.add_table(rows=int(len(lines) / 2) + 1, cols=2)

        while i < len(lines):
            if i % 2 == 0:
                Rightcell = table.cell(row, 0)
                Rightcell.text = lines[i].replace('\n', '')
                i = i + 1
            elif i % 2 == 1:
                LeftCell = table.cell(row, 1)
                LeftCell.text = lines[i].replace('\n', '')
                i = i + 1
                row += 1
            else:
                i = i + 1
                continue

        for row in table.rows:  #设置中文字体
            for paragraph in row.cells[1].paragraphs:
                paragraph.paragraph_format.first_line_indent = Pt(18)
                for run in paragraph.runs:
                    run.font.name = '华文楷体'
                    run.font.size = Pt(9)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')

            for paragraph in row.cells[0].paragraphs:  # 设置日文字体
                paragraph.paragraph_format.first_line_indent = Pt(9)
                for run in paragraph.runs:
                    run.font.name = 'MS Mincho'
                    run.font.size = Pt(9)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'MS Mincho')

    ProcessFile.close()
    document.save(".YiPai\.process\{}.docx".format(
        os.path.basename(file).replace(".md", "")))


# module\target\md2docx_1lang.py
def md2docx_1lang(file):
    document = Document()
    with open(file, 'r', encoding='UTF-8') as ProcessFile:
        lines = ProcessFile.readlines()
        i = 0
        row = 0
        table = document.add_table(rows=len(lines), cols=2)

        while i < len(lines):
            LeftCell = table.cell(row, 0)
            LeftCell.text = lines[i].replace('\n', '')
            i = i + 1
            row += 1

        for row in table.rows:
            for paragraph in row.cells[0].paragraphs:
                paragraph.paragraph_format.first_line_indent = Pt(9)
                for run in paragraph.runs:
                    run.font.name = 'MS Mincho'
                    run.font.size = Pt(9)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'MS Mincho')
    ProcessFile.close()
    document.save(".YiPai\.process\{}.docx".format(
        os.path.basename(file).replace(".md", "")))


# module\target\printdocx2pdf.py
def printdocx2pdf(file):
    InputFileName = os.path.basename(file)
    ProcessFileName = r".YiPai\.process\{}".format(InputFileName).replace(
        "docx", "pdf")
    ProcessFile = open(ProcessFileName, 'w')
    ProcessFile.close()
    convert(file, ProcessFileName
            )  # 注意file的数据类型是pathlib.WindowsPath，而ProcessFileName是字符串


# module\target\md_split.py
def Split(file):
    with open(file, 'r', encoding='UTF-8') as ProcessFile:
        FileTxt = ProcessFile.read()
        ProcessFile.close()
        InputLines = FileTxt.split('\n')
        wordcount = 0
        SplitStand = 800  # 分割标准
        SplitTime = 1  # 分割次数
        SplitCount = 800  # 分割时的字数，初始值
        OutputLines = []
        for InputLine in InputLines:
            wordcount += len(InputLine)

            if InputLine == '':
                continue

            if wordcount > SplitCount:  # 比较字数和分割标准
                OutputLine = InputLine + "\n\n# " + str(SplitCount)
                SplitTime += 1
                SplitCount = SplitStand * SplitTime  # 计算下一次分割标准
            else:
                OutputLine = InputLine
            OutputLines.append(OutputLine + "\n\n")
    with open(file, 'w', encoding='UTF-8') as OutPutFile:
        OutPutFile.writelines(OutputLines)
    OutPutFile.close()


# Action


def All2Half_(AllWideNums):
    HalfWideNumsList = []
    for num in AllWideNums:
        if ord(num) == 65292:  # 全角逗号
            HalfWideNumsList.append(num)
        elif ord(num) == 65311:  # 全角问号
            HalfWideNumsList.append(num)
        elif 65281 <= ord(num) <= 65374:
            HalfWideNum = chr(int(ord(num) - 65248))
            HalfWideNumsList.append(HalfWideNum)
        else:
            HalfWideNumsList.append(num)
        HalfWideNums = " ".join(map(str, HalfWideNumsList))
    return HalfWideNums


## 删除不必要的换行符


def del_slash(inputline):
    if "。\n" in inputline:
        outputline = inputline
    elif "」\n" in inputline:
        outputline = inputline
    elif "”\n" in inputline:
        outputline = inputline
    elif "？" in inputline:
        outputline = inputline
    else:
        outputline = inputline.replace("\n", "")
    return outputline


# action\FormatMarkdown.py


def formatmd(file):
    with open(file, 'r', encoding='UTF-8') as ProcessFile:
        OutputLines = []
        InputLines = ProcessFile.readlines()
        for InputLine in InputLines:
            if InputLine == "\n":  # 删掉空行
                continue
            else:
                OutputLine = re.sub(r"▼$", "。", InputLine)  # 删掉天声人语
                OutputLine = re.sub(r"^\s", "", OutputLine)
                OutputLine = All2Half_(OutputLine)
                OutputLine = OutputLine.replace(" ", "")
                OutputLine = del_slash(OutputLine)
                OutputLines.append(OutputLine)
    ProcessFile.close()
    with open(file, 'w', encoding='UTF-8') as OutPutFile:
        OutPutFile.writelines(OutputLines)
    OutPutFile.close()


# Result

## result\moveoutput.py


def moveoutput(file):
    oldname = os.path.abspath(file)
    newname = os.path.abspath(oldname).replace(".YiPai\.process", "")
    shutil.move(oldname, newname)


# Situation


# module\situation\docx2md.py
def Situation():
    docx2md_FileList = list(p.glob("**/*- 副本.docx"))
    for docx2md_File in docx2md_FileList:
        docx2md(docx2md_File)


# module\situation\md2md.py
def MDSituation():
    md2md_FileList = list(p.glob("**/*- 副本.md"))
    for md2md_File in md2md_FileList:
        print(md2md_File)
        md2md(md2md_File)


# Action
def Action():
    p = Path(ProcessFilePath)
    FileList = list(p.glob("**/*.md"))
    for file in FileList:
        formatmd(file)


# Target
def Target1lang():
    p = Path(ProcessFilePath)
    FileList = list(p.glob("**/*.md"))
    for file in FileList:
        md2docx_1lang(file)


def Target2lang():
    p = Path(ProcessFilePath)
    FileList = list(p.glob("**/*.md"))
    for file in FileList:
        md2docx_2lang(file)


def Print():
    p = Path(ProcessFilePath)
    FileList = list(p.glob("**/*.docx"))
    for file in FileList:
        printdocx2pdf(file)


# module\target\md_split.py
def TargetMD():
    p = Path(ProcessFilePath)
    FileList = list(p.glob("**/*.md"))
    for file in FileList:
        Split(file)


# Result
def Result():
    p = Path(ProcessFilePath)
    FileList = list(p.glob("**/*"))
    for file in FileList:
        # 如果勾选了MD那么就必须
        moveoutput(file)


def Run1lang(event):
    Situation()
    Action()
    Target1lang()
    Result()


def Run2lang(event):
    Situation()
    Action()
    Target2lang()
    Print()
    Result()


def RunMD(event):
    MDSituation()
    Action()
    TargetMD()
    Result()


class App:

    def __init__(self, master):
        frame = ttk.Frame(master)
        frame.pack()

        self.style = ttk.Style(frame)
        self.style.configure("TLabel",
                             foreground='#333333',
                             background="#ffffff")
        self.style.configure("TButton",
                             foreground='#333333',
                             background="#ffffff")
        self.style.configure("content.TFrame",
                             foreground='#333333',
                             background="#ffffff")
        self.style.configure('RunAppButton.TLabel',
                             foreground='#333333',
                             background="#ffffff",
                             sticky="nsew")

        self.content = ttk.Notebook(frame)
        self.content.pack()

        # Situation
        self.frame1 = ttk.Frame(self.content)
        self.content.add(self.frame1, text="Situation")
        InputFileType = tk.StringVar()
        self.menu_button = ttk.Menubutton(self.frame1, text='待处理文件格式')
        menu = tk.Menu(self.menu_button, tearoff=False)
        menu.add_radiobutton(label='纯文本Docx',
                             value='color',
                             variable=InputFileType)
        menu.add_radiobutton(label='TXT',
                             value='color',
                             variable=InputFileType)
        menu.add_radiobutton(label='双列表格Docx',
                             value='color',
                             variable=InputFileType)
        self.menu_button["menu"] = menu
        self.menu_button.grid(row=0, column=0)

        InputFileLanguage = tk.StringVar()
        self.menu_button = ttk.Menubutton(self.frame1, text='语言数')
        menu = tk.Menu(self.menu_button, tearoff=False)
        menu.add_radiobutton(label='1', value='1', variable=InputFileLanguage)
        menu.add_radiobutton(label='2', value='2', variable=InputFileLanguage)
        self.menu_button["menu"] = menu
        self.menu_button.grid(row=1, column=0)

        # Target
        self.frame2 = ttk.Frame(self.content)
        self.content.add(self.frame2, text="Target")

        self.SettingLabel = ttk.Label(self.frame2, text='字体', justify="center")
        self.SettingLabel.grid(row=1, column=0)
        self.SettingLabel = ttk.Label(self.frame2, text='字号', justify="center")
        self.SettingLabel.grid(row=2, column=0)
        self.SettingLabel = ttk.Label(self.frame2, text='缩进', justify="center")
        self.SettingLabel.grid(row=3, column=0)
        self.SettingLabel = ttk.Label(self.frame2, text='左', justify="center")
        self.SettingLabel.grid(row=0, column=1)
        self.SettingLabel = ttk.Label(self.frame2, text='右', justify="center")
        self.SettingLabel.grid(row=0, column=2)

        # 设定左列字体
        self.SetOutputFileFontName = ttk.Entry(
            self.frame2,
            justify="center",
        )
        self.SetOutputFileFontName.insert(0, "华文楷体")
        self.SetOutputFileFontName.grid(row=1, column=1)

        # 设定左列字号
        self.SetLeftFontsize = ttk.Entry(
            self.frame2,
            justify="center",
        )
        self.SetLeftFontsize.insert(0, "11")
        self.SetLeftFontsize.grid(row=2, column=1, sticky="nsew")

        # 设定左列缩进
        self.SetLeftFontsize = ttk.Entry(
            self.frame2,
            justify="center",
        )
        self.SetLeftFontsize.insert(0, "11")
        self.SetLeftFontsize.grid(row=3, column=1, sticky="nsew")

        # 设定右列字体
        self.SetFont = ttk.Entry(
            self.frame2,
            justify="center",
        )
        self.SetFont.insert(0, "MS Mincho")
        self.SetFont.grid(row=1, column=2)

        # 设定右列字号
        self.SetLeftFontsize = ttk.Entry(
            self.frame2,
            justify="center",
        )
        self.SetLeftFontsize.insert(0, "11")
        self.SetLeftFontsize.grid(row=2, column=2, sticky="nsew")
        # 设定右列缩进
        self.SetLeftFontsize = ttk.Entry(
            self.frame2,
            justify="center",
        )
        self.SetLeftFontsize.insert(0, "22")
        self.SetLeftFontsize.grid(row=3, column=2, sticky="nsew")

        self.SettingLabel = ttk.Label(self.frame2,
                                      text='导出格式',
                                      justify="center")
        self.SettingLabel.grid(row=4, column=0)
        self.SettingLabel = ttk.Label(self.frame2,
                                      text='docx',
                                      justify="center")
        self.SettingLabel.grid(row=5, column=0)
        self.SettingLabel = ttk.Label(self.frame2,
                                      text='txt/md',
                                      justify="center")
        self.SettingLabel.grid(row=6, column=0)
        self.SettingLabel = ttk.Label(self.frame2,
                                      text='pdf',
                                      justify="center")
        self.SettingLabel.grid(row=7, column=0)

        agreement = tk.StringVar()
        check_1 = ttk.Checkbutton(
            self.frame2,
            text="双列表格",
            variable=agreement,
            onvalue='agree',
            offvalue='disagree',
            # command=printinput,
        )
        check_1.grid(row=5, column=1, padx=5, pady=10, sticky="nsew")

        agreement = tk.StringVar()
        check_1 = ttk.Checkbutton(
            self.frame2,
            text="单列表格",
            variable=agreement,
            onvalue='agree',
            offvalue='disagree',
            # command=printinput,
        )
        check_1.grid(row=5, column=2, padx=5, pady=10, sticky="nsew")

        agreement = tk.StringVar()
        check_1 = ttk.Checkbutton(
            self.frame2,
            text="不分列",
            variable=agreement,
            onvalue='agree',
            offvalue='disagree',
            # command=printinput,
        )
        check_1.grid(row=5, column=3, padx=5, pady=10, sticky="nsew")

        agreement = tk.StringVar()
        check_1 = ttk.Checkbutton(
            self.frame2,
            text="TXT",
            variable=agreement,
            onvalue='agree',
            offvalue='disagree',
            # command=printinput,
        )
        check_1.grid(row=6, column=1, padx=5, pady=10, sticky="nsew")

        agreement = tk.StringVar()
        check_1 = ttk.Checkbutton(
            self.frame2,
            text="Markdown",
            variable=agreement,
            onvalue='agree',
            offvalue='disagree',
            # command=printinput,
        )
        check_1.grid(row=6, column=2, padx=5, pady=10, sticky="nsew")

        agreement = tk.StringVar()
        check_1 = ttk.Checkbutton(
            self.frame2,
            text="PDF",
            variable=agreement,
            onvalue='agree',
            offvalue='disagree',
            # command=printinput,
        )
        check_1.grid(row=7, column=1, padx=5, pady=10, sticky="nsew")
        # Action
        self.frame3 = ttk.Frame(self.content)
        self.content.add(self.frame3, text="Action")

        agreement = tk.StringVar()
        check_1 = ttk.Checkbutton(
            self.frame3,
            text="全半角数字转换",
            variable=agreement,
            onvalue='agree',
            offvalue='disagree',
            # command=printinput,
        )
        check_1.grid(row=0, column=0, padx=5, pady=10, sticky="nsew")
        self.content.add(self.frame3, text="Action")
        agreement = tk.StringVar()
        check_1 = ttk.Checkbutton(
            self.frame3,
            text="删除西文空格",
            variable=agreement,
            onvalue='agree',
            offvalue='disagree',
            # command=printinput,
        )
        check_1.grid(row=1, column=0, padx=5, pady=10, sticky="nsew")

        self.content.add(self.frame3, text="Action")
        agreement = tk.StringVar()
        check_1 = ttk.Checkbutton(
            self.frame3,
            text="空行",
            variable=agreement,
            onvalue='agree',
            offvalue='disagree',
            # command=printinput,
        )
        check_1.grid(row=2, column=0, padx=5, pady=10, sticky="nsew")

        self.content.add(self.frame3, text="Action")
        agreement = tk.StringVar()
        check_1 = ttk.Checkbutton(
            self.frame3,
            text="《天声人语》段尾标记替换为句号",
            variable=agreement,
            onvalue='agree',
            offvalue='disagree',
            # command=printinput,
        )
        check_1.grid(row=3, column=0, padx=5, pady=10, sticky="nsew")
        # model_result
        self.frame4 = ttk.Frame(self.content)
        self.content.add(self.frame4, text="Result")

        self.frame5 = ttk.Frame(self.content)
        self.content.add(self.frame5, text="Run")
        self.RunAppButton = ttk.Button(
            self.frame5, text="单语文件转换")  #,style='RunAppButton.TLabel'
        self.RunAppButton.bind("<Button-1>", Run1lang)
        self.RunAppButton.grid(row=0, column=0)

        self.RunAppButton = ttk.Button(
            self.frame5, text="双语文件转换")  #,style='RunAppButton.TLabel'
        self.RunAppButton.bind("<Button-1>", Run2lang)
        self.RunAppButton.grid(row=1, column=0)

        self.RunAppButton = ttk.Button(
            self.frame5, text="MD格式转换")  #,style='RunAppButton.TLabel'
        self.RunAppButton.bind("<Button-1>", RunMD)
        self.RunAppButton.grid(row=2, column=0)


if __name__ == "__main__":
    root = tk.Tk()
    root.title("YiPai")

    ctypes.windll.shcore.SetProcessDpiAwareness(1)
    ScaleFactor = ctypes.windll.shcore.GetScaleFactorForDevice(0)
    root.iconphoto(True, tk.PhotoImage(file='.YiPai\main.png'))
    root.call('tk', 'scaling', ScaleFactor / 75) #
    root.resizable(False, False)

    root.tk.call("source", r".YiPai\sun-valley.tcl")
    root.tk.call("set_theme", "light")
    app = App(root)
    root.attributes("-topmost", True)
    root.mainloop()
